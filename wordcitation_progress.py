import tkinter as tk
from tkinter import ttk
from tkinter import filedialog, messagebox
from docx import Document
from ttkthemes import ThemedTk
import re
import os

# 函数定义：

def extract_references(doc_content):
    pattern = r"\[(\d+)\]\s*(.+)"
    matches = re.findall(pattern, doc_content)
    return {num.strip(): title.strip() for num, title in matches}

def find_citekeys_fuzzy(bib_content, references):
    citekeys = {}
    for num, title in references.items():
        bib_entries = re.findall(r"@.+{(.*?),\s+title={(.+?)}", bib_content)
        for citekey, bib_title in bib_entries:
            if title.lower() in bib_title.lower() or bib_title.lower() in title.lower():
                citekeys[num] = citekey
                break
    return citekeys

def replace_references_in_text(text, citekeys):
    # Use a regex that matches the reference pattern, including composite references and surrounding brackets
    pattern = re.compile(r'\[\s*([\d,\s]+)\s*\]')

    # Define a replacement function that will be used by re.sub
    def reference_replacer(match):
        # Extract the individual reference numbers from the match
        refs = match.group(1).replace(' ', '').split(',')
        # Replace each reference number with the corresponding citekey if it exists
        replaced_refs = [f"\\cite{{{citekeys[ref]}}}" if ref in citekeys else ref for ref in refs]
        # Join the replaced references back into a composite reference if necessary
        return ', '.join(replaced_refs) if len(replaced_refs) > 1 else replaced_refs[0]

    # Replace the references in the text using the replacement function
    updated_text = pattern.sub(reference_replacer, text)
    return updated_text


def replace_references_in_text(text, citekeys):
    # Define a regex pattern that matches both individual and composite references
    pattern = re.compile(r'\[\s*([\d,\s]+)\s*\]')

    # Define a replacement function for re.sub
    def reference_replacer(match):
        # Extract individual reference numbers from the match
        refs = match.group(1).replace(' ', '').split(',')
        # Replace each reference number with the corresponding citekey if it exists
        replaced_refs = [f"\\cite{{{citekeys.get(ref, ref)}}}" for ref in refs]
        # Join the replaced references back into a composite reference if necessary
        return ', '.join(replaced_refs)

    # Replace the references in the text using the replacement function
    updated_text = pattern.sub(reference_replacer, text)
    return updated_text

def process_documents(doc1_path, doc2_path, bib_path, save_path):
    try:
        progress_var.set(0)  # Initialize progress
        root.update()

        # Read Document 1
        doc1 = Document(doc1_path)
        doc1_content = '\n'.join([para.text for para in doc1.paragraphs])
        progress_var.set(10)  # Update progress
        root.update()

        # Read Document 2
        doc2 = Document(doc2_path)
        doc2_content = '\n'.join([para.text for para in doc2.paragraphs])
        progress_var.set(20)  # Update progress
        root.update()

        # Read .bib file
        with open(bib_path, 'r', encoding='utf-8') as file:
            bib_content = file.read()
        progress_var.set(30)  # Update progress
        root.update()

        # Extract and find references
        references_doc2 = extract_references(doc2_content)
        citekeys = find_citekeys_fuzzy(bib_content, references_doc2)
        progress_var.set(50)  # Update progress
        root.update()

        # Replace references in Document 1
        updated_content = replace_references_in_text(doc1_content, citekeys)
        progress_var.set(70)  # Update progress
        root.update()

        # Create and save the new document
        new_doc = Document()
        for paragraph_text in updated_content.split('\n'):
            new_doc.add_paragraph(paragraph_text)
        new_doc.save(save_path)
        progress_var.set(80)  # Update progress
        root.update()

        # Check for missing references
        refs_doc1 = set(re.findall(r'\[\s*(\d+)\s*\]', doc1_content))
        missing_in_doc1 = refs_doc1 - set(references_doc2.keys())
        missing_in_bib = set(references_doc2.keys()) - set(citekeys.keys())

        # Save missing references notice if needed
        if missing_in_doc1 or missing_in_bib:
            missing_notice = ""
            if missing_in_doc1:
                missing_notice += "References in Document 1 not found in Document 2:\n" + ', '.join(missing_in_doc1) + '\n\n'
            if missing_in_bib:
                missing_notice += "References in Document 2 not found in .bib:\n" + ', '.join(missing_in_bib)
            notice_path = os.path.join(os.path.dirname(save_path), "missing_references.txt")
            with open(notice_path, 'w', encoding='utf-8') as file:
                file.write(missing_notice)
            messagebox.showinfo("Missing References", "The missing references are saved to missing_reference.txt in the same location as the output document.")

        messagebox.showinfo("Success", "Document processed and saved successfully.")
        progress_var.set(100)  # Finalize progress
        root.update()

    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {e}")
        progress_var.set(0)  # Reset progress on error
        root.update()

def select_file1():
    path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx"), ("All files", "*.*")])
    if path:
        doc1_path.set(path)

def select_file2():
    path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx"), ("All files", "*.*")])
    if path:
        doc2_path.set(path)

def select_bib():
    path = filedialog.askopenfilename(filetypes=[("Bib files", "*.bib"), ("All files", "*.*")])
    if path:
        bib_path.set(path)

def select_save_path():
    path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx"), ("All files", "*.*")])
    if path:
        new_doc_path.set(path)

# https://ttkthemes.readthedocs.io/en/latest/themes.html#yaru
# Create the themed main window
# Initialize the themed Tk window
root = ThemedTk(theme="arc")
root.title("Reference Formatter")
root.geometry("600x300")  # Adjust as needed

# Create frames
frame1 = ttk.Frame(root, padding="10")
frame1.pack(fill='x')
frame2 = ttk.Frame(root, padding="10")
frame2.pack(fill='x')
frame3 = ttk.Frame(root, padding="10")
frame3.pack(fill='x')
frame4 = ttk.Frame(root, padding="10")
frame4.pack(fill='x')
frame5 = ttk.Frame(root, padding="10")
frame5.pack(fill='x')

# Define path variables
doc1_path = tk.StringVar()
doc2_path = tk.StringVar()
bib_path = tk.StringVar()
new_doc_path = tk.StringVar()

# Create and place widgets with consistent width
entry_width = 50  # You may need to adjust this based on the window size

# Document 1 widgets
ttk.Label(frame1, text="Content (.docx):").pack(side='left')
doc1_entry = ttk.Entry(frame1, textvariable=doc1_path, width=entry_width)
doc1_entry.pack(side='left', fill='x', expand=True)
ttk.Button(frame1, text="Browse", command=select_file1).pack(side='right')

# Document 2 widgets
ttk.Label(frame2, text="Reference List (.docx):").pack(side='left')
doc2_entry = ttk.Entry(frame2, textvariable=doc2_path, width=entry_width)
doc2_entry.pack(side='left', fill='x', expand=True)
ttk.Button(frame2, text="Browse", command=select_file2).pack(side='right')

# .bib File widgets
ttk.Label(frame3, text="Bibtex File (.bib):").pack(side='left')
bib_entry = ttk.Entry(frame3, textvariable=bib_path, width=entry_width)
bib_entry.pack(side='left', fill='x', expand=True)
ttk.Button(frame3, text="Browse", command=select_bib).pack(side='right')

# Save path widgets
ttk.Label(frame4, text="Save updated document as:").pack(side='left')
save_entry = ttk.Entry(frame4, textvariable=new_doc_path, width=entry_width)
save_entry.pack(side='left', fill='x', expand=True)
ttk.Button(frame4, text="Save As", command=select_save_path).pack(side='right')

# Progress Bar frame
progress_frame = ttk.Frame(root, padding="10")
progress_frame.pack(fill='x')

# Progress Bar widget
progress_var = tk.DoubleVar()  # Variable to track progress
progress_bar = ttk.Progressbar(progress_frame, variable=progress_var, maximum=100, length=250)
progress_bar.pack()

# Process button frame
process_frame = ttk.Frame(root, padding="10")
process_frame.pack(fill='x')

# Shorter process button with specific style
process_button = ttk.Button(
    process_frame,
    text="Process Documents",
    command=lambda: process_documents(doc1_entry.get(), doc2_entry.get(), bib_entry.get(), new_doc_path.get())  # Include new_doc_path
)
process_button.pack()

# Run the main loop
root.mainloop()